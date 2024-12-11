import aspose.words as aw
import docx
from docx import *
from docx.text.paragraph import Paragraph
import xml.etree.ElementTree as ET
from docx.document import Document as doctwo
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import pandas as pd
from xml.etree import ElementTree
from io import StringIO
import io
import csv
import base64
import pandas as pd


#This function extracts the tables and paragraphs from the document object
def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, doctwo):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


#This function extracts the table from the document object as a dataframe
def read_docx_tables(document, tab_id=None, **kwargs):
    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in document.tables]
    else:
        try:
            return read_docx_tab(document.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise

def list_to_markdown(element):
    title, items = element[0], element[1]

    # Start with the title
    markdown = f"{title}\n\n"

    # Add each item as a markdown list entry
    for item in items:
        markdown += f"- {item}\n"

    return markdown

def table_to_markdown(element):
    title, table = element[0], element[1]

    markdown = f"{title}\n\n"

    table_md = table.to_markdown(index=False)
    markdown += f"{table_md}\n"

    return markdown

def get_docx_content(document):
    combined_df = pd.DataFrame(columns=['para_text','table_id','style'])
    table_mod = pd.DataFrame(columns=['string_value','table_id'])

    image_df = pd.DataFrame(columns=['image_index','image_rID','image_filename','image_base64_string'])

    table_list=[]
    xml_list=[]

    i=0
    imagecounter = 0


    blockxmlstring = ''
    for block in iter_block_items(document):
        if 'text' in str(block):
            isappend = False

            runboldtext = ''
            for run in block.runs:
                if run.bold:
                    runboldtext = runboldtext + run.text

            style = str(block.style.name)

            if '<w:numId' in str(block._p.xml):
                style = 'List Paragraph'
            
            appendtxt = str(block.text)
            appendtxt = appendtxt.replace("\n","")
            appendtxt = appendtxt.replace("\r","")
            tabid = 'Novalue'
            paragraph_split = appendtxt.lower().split()

            isappend = True
            for run in block.runs:
                xmlstr = str(run.element.xml)
                my_namespaces = dict([node for _, node in ElementTree.iterparse(StringIO(xmlstr), events=['start-ns'])])
                root = ET.fromstring(xmlstr)
                #Check if pic is there in the xml of the element. If yes, then extract the image data
                if 'pic:pic' in xmlstr:
                    xml_list.append(xmlstr)
                    for pic in root.findall('.//pic:pic', my_namespaces):
                        cNvPr_elem = pic.find("pic:nvPicPr/pic:cNvPr", my_namespaces)
                        name_attr = cNvPr_elem.get("name")
                        blip_elem = pic.find("pic:blipFill/a:blip", my_namespaces)
                        embed_attr = blip_elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        isappend = True
                        appendtxt = str('Document_Imagefile/' + name_attr + '/' + embed_attr + '/' + str(imagecounter))
                        document_part = document.part
                        image_part = document_part.related_parts[embed_attr]
                        image_base64 = base64.b64encode(image_part._blob)
                        image_base64 = image_base64.decode()
                        dftemp = pd.DataFrame({'image_index':[imagecounter],'image_rID':[embed_attr],'image_filename':[name_attr],'image_base64_string':[image_base64]})
                        image_df = pd.concat([image_df,dftemp],sort=False)
                        style = 'Novalue'
                    imagecounter = imagecounter + 1

        elif 'table' in str(block):
            isappend = True
            style = 'Novalue'
            appendtxt = str(block)
            tabid = i
            dfs = read_docx_tables(tab_id=i)
            dftemp = pd.DataFrame({'para_text':[appendtxt],'table_id':[i],'style':[style]})
            table_mod = pd.concat([table_mod,dftemp],sort=False)
            table_list.append(dfs)
            i=i+1
        if isappend:
                dftemp = pd.DataFrame({'para_text':[appendtxt],'table_id':[tabid],'style':[style]})
                combined_df = pd.concat([combined_df,dftemp],sort=False)

    combined_df = combined_df.reset_index(drop=True)
    image_df = image_df.reset_index(drop=True)

    for table in table_list:
        for col in table.columns:
            table[col] = table[col].apply(lambda x: f"{col} {x}")
    
    return combined_df, table_list, image_df

def docx2md(file_path):
    doc = aw.Document(file_path)
    doc.save("data\docx2md.md")
    with open("data\docx2md.md", "r", encoding='utf-8') as file:
        doc = file.read()

    document = doc[248:-85]
    return document

def docx2txt(file_path):
    doc = docx.Document(file_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def content_to_markdown(document):
    combined_df, table_list, image_df = get_docx_content(document)
    #with pd.option_context('display.max_rows', None, 'display.max_columns', None):
        #print(combined_df)
    chunks = []
    current_list_title = None
    current_list_items = []

    index = 0
    while index < len(combined_df):
        #print(index)
        row = combined_df.iloc[index]
        next_row = combined_df.iloc[index + 1] if index + 1 < len(combined_df) else None

        if row['table_id'] == 'Novalue':
            if index + 1 < len(combined_df) and next_row['style'] == 'List Paragraph':
                current_list_title = row['para_text']
                current_list_items = []
                index += 1
                next_row = combined_df.iloc[index]

                while index < len(combined_df) and next_row['style'] == 'List Paragraph':
                    current_list_items.append(next_row['para_text'])
                    index += 1
                    next_row = combined_df.iloc[index] if index < len(combined_df) else None

                filtered_title = current_list_title
                chunks.append(('list', (filtered_title, current_list_items)))
                current_list_title = None
                current_list_items = []
            elif index + 1 < len(combined_df) and next_row['style'] == 'Normal':
                # sentences = re.split(r'\.(?=\s|$)', row['para_text'])
                # for sentence in sentences:
                #     sentence = sentence.strip()
                #     if sentence:
                #         chunks.append(('sentence', sentence))
                if (row['para_text'] != ''):
                    chunks.append(('paragraph', row['para_text']))
                index += 1
            # else:
            #     index += 1
            elif index + 1 < len(combined_df) and next_row['table_id'] != 'Novalue':
                table_title = row['para_text']
                table_ref = next_row['para_text']
                table = table_list[next_row['table_id']]
                chunks.append(('table', (table_title, table)))
                index += 2
            else:
                index += 1
        else:
            index += 1

    chunk_md = []
    chunks_type = []

    for chunk_type, chunk in chunks:
        chunks_type.append(chunk_type)
        if chunk_type == 'list':
            chunk_md.append(list_to_markdown(chunk))
        elif chunk_type == 'paragraph':
            chunk_md.append(chunk)
        elif chunk_type == 'table':
            chunk_md.append(table_to_markdown(chunk))

    return chunk_md, chunks_type

